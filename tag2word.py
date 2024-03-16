from docx import Document
from pdf2image import convert_from_path
from docx.shared import Cm
import argparse
import os

MIN_DPI = 300
MAX_DPI = 1200
ECO_MODE_ENABLE = 1
ECO_MODE_DISABLE = 0
A4_WIDTH = float(21.0)
A4_HEIGHT = float(29.7)
TRIMAN_PATH = "./images/triman.png"

def pdf2png(pdf_path, output_path, image_dpi):
    pages = convert_from_path(pdf_path, dpi=image_dpi)
    i=0
    file_names = []
    insert_num = []
    for page in pages:
        i+=1
        file_name = output_path + '/out' + str(i) + '.png'
        file_names.append(file_name)
        insert_num.append(1)
        page.save(file_name, 'PNG')
    return file_names, insert_num

def insert_images(doc, image_paths, insert_counts, image_width, image_height, image_dpi):
    """
    插入图片

    Args:
      doc: 文档对象
      image_paths: 图片路径列表
      insert_counts: 插入张数列表

    Returns:
      None
    """
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    for image_path, insert_count in zip(image_paths, insert_counts):
        for _ in range(insert_count):
            p = run.add_picture(image_path, width=Cm(image_width), height=Cm(image_height))

def main():
    """
    主函数
    """
    # 创建参数解析器
    parser = argparse.ArgumentParser()

    # 添加参数
    parser.add_argument('-p', '--pdf_path', required=True, help='pdf文件路径')
    parser.add_argument('-d', '--dpi', type=int, default=300, help='图片清晰度dpi')
    parser.add_argument('-o', '--output_dir', required=True, help='输出文件目录')
    parser.add_argument('-e', '--eco_mode', type=int, default=1, help='输出文件目录')

    # 解析参数
    args = parser.parse_args()

    # 获取参数
    pdf_path = args.pdf_path
    image_dpi = args.dpi
    outpu_dir = args.output_dir
    if image_dpi < MIN_DPI:
        image_dpi = MIN_DPI
    if image_dpi > MAX_DPI:
        image_dpi = MAX_DPI
    eco_mode = args.eco_mode
    if eco_mode > 0:
        eco_mode = ECO_MODE_ENABLE
    if eco_mode <= 0:
        eco_mode = ECO_MODE_DISABLE
    paper_width = A4_WIDTH
    paper_height = A4_HEIGHT
    # 单位都是厘米
    left_margin = float(0.3) 
    right_margin = float(0.3)
    top_margin = float(1)
    bottom_margin = float(1)

    image_width = float(6.75)
    image_height = float(2)

    # PDF文件转PNG
    pdf2image_path = outpu_dir + "/"
    if not os.path.isdir(pdf2image_path):
        os.makedirs(pdf2image_path)
    file_names, insert_num = pdf2png(pdf_path, pdf2image_path, image_dpi)

    # 将一页中剩余的部分打印triman标签
    if eco_mode:
        column_num = int((paper_width - left_margin - right_margin) / image_width)
        row_num = int((paper_height - top_margin - bottom_margin) / image_height)
        img_num_per_page = column_num * row_num
        img_num_left = img_num_per_page - sum(insert_num) % img_num_per_page
        print(img_num_per_page, sum(insert_num), img_num_left)
        file_names.append(TRIMAN_PATH)
        insert_num.append(img_num_left)


    # 创建一个新的word文档
    doc = Document()
    # 设置纸张大小
    section = doc.sections[0]
    section.page_width = Cm(paper_width)
    section.page_height = Cm(paper_height)
    
    # 设置页边距
    section.left_margin = Cm(left_margin)
    section.right_margin = Cm(right_margin)
    section.top_margin = Cm(top_margin)
    section.bottom_margin = Cm(bottom_margin)

    # 设置行间距
    paragraph_format = doc.styles['Normal'].paragraph_format
    paragraph_format.line_spacing = 1
    print(f'line_spacing = {paragraph_format.line_spacing}')


    # 插入图片
    insert_images(doc, file_names, insert_num, image_width, image_height, image_dpi)
    # 保存文档
    doc.save(f'{pdf2image_path}/output.docx')

if __name__ == '__main__':
  main()

